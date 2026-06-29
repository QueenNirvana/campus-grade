from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SOURCE = Path("D:/15Pro/Desktop/高校成绩管理系统课程设计报告.docx")
OUT = Path("D:/15Pro/Desktop/高校成绩管理系统课程设计报告-按模板修订版.docx")
REFERENCE_COVER = Path("D:/15Pro/Desktop/2404 张欢宁，赵文豪JAVA程序设计-课设报告(1).docx")


def set_run_font(run, size=11, bold=False, name="宋体", ascii_name="Times New Roman", color=None):
    run.font.name = ascii_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=11, bold=False, name="宋体", ascii_name="Times New Roman", color=None):
    for run in paragraph.runs:
        set_run_font(run, size, bold, name, ascii_name, color)


def add_center_text(doc, text, size=16, bold=False, before=0, after=0, name="宋体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, name=name)
    return p


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    if p.runs:
        p.runs[0].text = text
        run = p.runs[0]
    else:
        run = p.add_run(text)
    set_run_font(run, size=12, bold=bold)


def build_cover_elements(doc):
    ref_doc = Document(str(REFERENCE_COVER))
    values = [
        "Java课程设计",
        "物联网2403",
        "周成林 逯雨喆",
        "241040700321 241040700308",
        "曹丹 夏娟",
        "2026.6.8—2026.6.20",
    ]
    for row, value in zip(ref_doc.tables[0].rows, values):
        cell = row.cells[1]
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = value
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.add_run(value)

    elements = []
    for element in list(ref_doc._body._element):
        text = "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()
        if text == "目   次":
            break
        elements.append(deepcopy(element))
    return elements


def remove_elements_before_toc(doc):
    body = doc._body._element
    for child in list(body):
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if text.strip() == "目录":
            break
        body.remove(child)


def insert_elements_at_start(doc, elements):
    body = doc._body._element
    first = body[0]
    for element in elements:
        if "__COVER_MARKER__" in "".join(t.text or "" for t in element.iter(qn("w:t"))):
            continue
        body.insert(body.index(first), element)


def replace_paragraph_text(doc, old, new):
    for p in doc.paragraphs:
        if p.text.strip() == old:
            p.clear()
            run = p.add_run(new)
            if new in {"一、绪论", "四、系统设计", "六、系统实现", "九、系统测试", "十二、总结与展望", "参考文献"}:
                set_run_font(run, size=16, bold=True, name="黑体", color="2E74B5")
            else:
                set_run_font(run, size=11)


def replace_inline_text(doc, old, new):
    for p in doc.paragraphs:
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)


def add_report_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_report_heading(doc, text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, name="黑体", color="1F4D78")
    return p


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code.strip())
    set_run_font(run, size=8.5, name="Consolas", ascii_name="Consolas")
    doc.add_paragraph()
    return table


def add_explain_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(["代码位置", "核心代码", "作用说明"]):
        set_cell(table.rows[0].cells[i], header, bold=True)
    for item in rows:
        cells = table.add_row().cells
        for i, value in enumerate(item):
            set_cell(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    return table


def insert_layered_code_explanation(doc):
    body = doc._body._element
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "九、系统测试":
            target = p._p
    if target is None:
        return

    scratch = Document()
    add_report_heading(scratch, "8.7 三层架构代码讲解：Controller、Service、Mapper", size=13)
    add_report_paragraph(
        scratch,
        "从系统实现角度看，本项目采用清晰的分层结构说明核心代码。持久层没有使用 JPA Repository，"
        "而是按照课程要求使用原生 MyBatis。Controller 层负责接收浏览器请求，Service 层负责业务规则和权限判断，"
        "Mapper 接口与 XML SQL 共同承担 Repository/DAO 层的数据访问职责。"
    )
    add_explain_table(scratch, [
        ("Controller 层", "@RestController、@RequestMapping、@GetMapping、@PostMapping", "定义 REST 接口，接收前端 Axios 请求，把请求参数封装为 DTO，并调用 Service。"),
        ("Service 层", "@Service、GradeService、GradeRuleService", "集中处理成绩保存、课程归属校验、课程权重计算、绩点区间匹配等核心业务。"),
        ("Mapper/Repository 层", "GradeMapper 接口、GradeMapper.xml", "承担持久层的数据访问职责，通过手写 SQL 完成查询、插入、更新、删除和统计。"),
        ("Entity/DTO 层", "Grade、Course、GradeSaveRequest、CourseGradeRuleConfig", "Entity 对应数据库业务对象，DTO 用于接收前端请求或返回配置数据。"),
    ])

    add_report_heading(scratch, "8.8 Controller 层：成绩接口接收与分发", size=13)
    add_report_paragraph(
        scratch,
        "Controller 是前后端交互入口。以 GradeController 为例，前端成绩管理页面访问 /api/grades，"
        "查询时调用 list 方法，新增成绩时调用 save 方法，修改时调用 update 方法，删除时调用 delete 方法。"
        "Controller 本身不直接计算成绩，也不直接写 SQL，而是把请求交给 GradeService 处理。这样可以保持接口层简洁，"
        "也便于把权限判断和成绩计算统一放在业务层。"
    )
    add_code_block(scratch, """
@RestController
@RequestMapping("/api/grades")
public class GradeController {
    private final GradeService gradeService;

    public GradeController(GradeService gradeService) {
        this.gradeService = gradeService;
    }

    @GetMapping
    public ApiResponse<List<Grade>> list() {
        return ApiResponse.ok(gradeService.visibleGrades());
    }

    @PostMapping
    public ApiResponse<Grade> save(@Valid @RequestBody GradeSaveRequest request) {
        return ApiResponse.ok(gradeService.save(request));
    }

    @PutMapping("/{id}")
    public ApiResponse<Grade> update(@PathVariable Long id,
                                     @Valid @RequestBody GradeSaveRequest request) {
        request.id = id;
        return ApiResponse.ok(gradeService.save(request));
    }

    @DeleteMapping("/{id}")
    public ApiResponse<Void> delete(@PathVariable Long id) {
        gradeService.delete(id);
        return ApiResponse.ok(null);
    }
}
""")
    add_report_paragraph(
        scratch,
        "上述代码中，@RestController 表示该类返回 JSON 数据；@RequestMapping(\"/api/grades\") 统一定义成绩接口路径；"
        "@RequestBody 把前端提交的 JSON 转换为 GradeSaveRequest 对象；@PathVariable 用于读取 URL 中的成绩 ID。"
        "所有返回值都包装为 ApiResponse，保证前端收到统一的 code、message、data 响应结构。"
    )

    add_report_heading(scratch, "8.9 Service 层：成绩业务规则集中处理", size=13)
    add_report_paragraph(
        scratch,
        "Service 层是本系统最重要的业务层。成绩保存时，系统不能信任前端直接传入 total_score 和 grade_point，"
        "而是必须在后端读取课程自己的 usual_weight、final_weight 和 course_grade_rules，再计算总评成绩与绩点。"
        "同时，教师只能操作本人课程，学生不能录入或修改成绩，这些限制也都放在 Service 层完成。"
    )
    add_code_block(scratch, """
public Grade save(GradeSaveRequest request) {
    if (SecurityUtils.isStudent()) {
        throw new BusinessException("学生不能录入或修改成绩");
    }
    courseService.assertTeacherOwnsCourse(request.courseId);
    Course course = courseService.findRequiredCourse(request.courseId);
    List<CourseGradeRule> rules = ruleMapper.findByCourseId(request.courseId);

    BigDecimal totalScore = gradeRuleService.calculateTotalScore(
            request.usualScore,
            request.finalScore,
            course.usualWeight,
            course.finalWeight
    );
    BigDecimal gradePoint = gradeRuleService.calculateGradePoint(rules, totalScore);

    Grade grade = new Grade();
    grade.studentId = request.studentId;
    grade.courseId = request.courseId;
    grade.usualScore = request.usualScore;
    grade.finalScore = request.finalScore;
    grade.totalScore = totalScore;
    grade.gradePoint = gradePoint;
    return saveOrUpdate(grade);
}
""")
    add_report_paragraph(
        scratch,
        "这段代码体现了成绩模块的核心流程：第一步判断当前用户是否为学生；第二步判断教师是否拥有该课程；"
        "第三步读取课程权重和绩点规则；第四步根据平时成绩、期末成绩、课程权重计算 total_score；"
        "第五步根据该课程自己的绩点区间计算 grade_point；最后再写入 grades 表。"
    )
    add_code_block(scratch, """
public BigDecimal calculateTotalScore(BigDecimal usualScore, BigDecimal finalScore,
                                      BigDecimal usualWeight, BigDecimal finalWeight) {
    validateScore(usualScore, "平时成绩");
    validateScore(finalScore, "期末成绩");
    validateWeights(usualWeight, finalWeight);
    return usualScore.multiply(usualWeight)
            .add(finalScore.multiply(finalWeight))
            .divide(new BigDecimal("100"), 2, RoundingMode.HALF_UP);
}

public void validateWeights(BigDecimal usualWeight, BigDecimal finalWeight) {
    if (usualWeight.add(finalWeight).compareTo(new BigDecimal("100")) != 0) {
        throw new BusinessException("平时成绩占比和期末成绩占比之和必须等于100");
    }
}
""")
    add_report_paragraph(
        scratch,
        "权重校验和总评成绩计算单独放在 GradeRuleService 中，便于复用和单元测试。"
        "如果某门课程设置为 30% 平时成绩和 70% 期末成绩，就按 30/70 计算；另一门课程设置为 40/60，则按 40/60 计算。"
        "这正好体现了本系统“成绩计算规则按课程配置，而不是全局统一”的设计重点。"
    )

    add_report_heading(scratch, "8.10 Mapper/Repository 层：MyBatis 数据访问实现", size=13)
    add_report_paragraph(
        scratch,
        "在常见三层结构中，持久层负责数据库访问。本项目采用 MyBatis，因此分成 Mapper 接口和 XML SQL 两部分。"
        "Mapper 接口声明 Java 方法，XML 文件编写实际 SQL。这样做的优点是 SQL 语句清楚可见，适合课程设计中展示数据库访问过程。"
    )
    add_code_block(scratch, """
public interface GradeMapper {
    List<Grade> findAll();
    List<Grade> findByTeacherId(Long teacherId);
    List<Grade> findByStudentId(Long studentId);
    Grade findById(Long id);
    int insert(Grade item);
    int update(Grade item);
    int delete(Long id);
    CourseStats courseStats(Long courseId);
    StudentSummary studentSummary(Long studentId);
}
""")
    add_report_paragraph(
        scratch,
        "GradeMapper 中的方法与成绩管理页面和统计页面一一对应。例如 findByTeacherId 用于教师查看本人课程成绩，"
        "findByStudentId 用于学生只查看本人课程成绩，courseStats 和 studentSummary 用于统计分析页面。"
    )
    add_code_block(scratch, """
<select id="findByTeacherId" resultType="com.example.campusgrade.entity.Grade">
    <include refid="gradeSelect"/>
    WHERE c.teacher_id=#{teacherId}
    ORDER BY g.id DESC
</select>

<select id="findByStudentId" resultType="com.example.campusgrade.entity.Grade">
    <include refid="gradeSelect"/>
    WHERE g.student_id=#{studentId}
    ORDER BY c.course_name
</select>

<insert id="insert" useGeneratedKeys="true" keyProperty="id">
    INSERT INTO grades(student_id, course_id, usual_score, final_score,
                       total_score, grade_point, remark)
    VALUES(#{studentId}, #{courseId}, #{usualScore}, #{finalScore},
           #{totalScore}, #{gradePoint}, #{remark})
</insert>
""")
    add_report_paragraph(
        scratch,
        "这部分 XML SQL 说明了持久层如何真正访问数据库。findByTeacherId 通过课程表中的 teacher_id 过滤，"
        "避免教师看到其他教师课程；findByStudentId 通过 student_id 过滤，保证学生只能查看本人记录；insert 语句写入 grades 表时，"
        "保存的不只是 usual_score 和 final_score，还包括 Service 层计算得到的 total_score 与 grade_point。"
    )
    add_report_heading(scratch, "8.11 分层调用流程总结", size=13)
    add_explain_table(scratch, [
        ("1. 前端页面", "Axios 调用 POST /api/grades", "页面只提交学生、课程、平时成绩、期末成绩和备注。"),
        ("2. Controller", "GradeController.save(request)", "接收 JSON 请求并转换为 GradeSaveRequest。"),
        ("3. Service", "GradeService.save(request)", "校验角色和课程归属，调用规则服务计算总评成绩和绩点。"),
        ("4. 规则服务", "GradeRuleService.calculateTotalScore / calculateGradePoint", "校验权重之和为 100，按课程规则计算 total_score 和 grade_point。"),
        ("5. Mapper", "GradeMapper.insert/update", "调用 XML SQL，把最终成绩记录写入 MySQL。"),
        ("6. 前端展示", "ApiResponse<Grade>", "返回保存后的成绩记录，表格刷新显示自动计算结果。"),
    ])

    insert_pos = body.index(target)
    for element in list(scratch._body._element):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(insert_pos, deepcopy(element))
        insert_pos += 1


def insert_reference_section(doc):
    refs = [
        "[1] Oracle. Java SE Documentation. Oracle Corporation.",
        "[2] Craig Walls. Spring in Action, Sixth Edition. Manning Publications, 2022.",
        "[3] MyBatis Team. MyBatis 3 User Guide.",
        "[4] MySQL 8.0 Reference Manual. Oracle Corporation.",
        "[5] Vue.js Team. Vue 3 Documentation.",
        "[6] Element Plus Team. Element Plus Component Documentation.",
    ]
    body = doc._body._element
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "附录：主要接口清单":
            target = p._p
    if target is None:
        return

    scratch = Document()
    heading = scratch.add_paragraph()
    heading_run = heading.add_run("参考文献")
    set_run_font(heading_run, size=16, bold=True, name="黑体", color="2E74B5")
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    for ref in refs:
        p = scratch.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ref)
        set_run_font(run, size=10.5)
    scratch.add_paragraph()

    insert_pos = body.index(target)
    for element in list(scratch._body._element):
        if element.tag.endswith("sectPr"):
            continue
        body.insert(insert_pos, deepcopy(element))
        insert_pos += 1


def update_contents(doc):
    desired = [
        "摘要",
        "一、绪论",
        "二、小组分工明细",
        "三、需求分析",
        "四、系统设计",
        "五、数据库设计",
        "六、系统实现",
        "七、程序流程图",
        "八、核心关键代码",
        "九、系统测试",
        "十、系统运行与部署说明",
        "十一、系统截图",
        "十二、总结与展望",
        "参考文献",
        "附录：主要接口清单",
    ]
    toc = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "目录":
            toc = i
            break
    if toc is None:
        return
    for offset, text in enumerate(desired, start=1):
        if toc + offset < len(doc.paragraphs):
            p = doc.paragraphs[toc + offset]
            p.clear()
            run = p.add_run(text)
            set_run_font(run, size=11)


def normalize_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if p.text.strip() == "目录":
            for run in p.runs:
                set_run_font(run, size=16, bold=True, name="黑体")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
        elif p.text.strip().startswith(("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、", "十一、", "十二、")) or p.text.strip() in {"摘要", "参考文献", "附录：主要接口清单", "课程设计任务书信息"}:
            for run in p.runs:
                set_run_font(run, size=15 if p.text.strip() != "摘要" else 16, bold=True, name="黑体", color="2E74B5")
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
        else:
            for run in p.runs:
                if run.text:
                    set_run_font(run, size=10.5)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(4)


def main():
    doc = Document(SOURCE)
    remove_elements_before_toc(doc)
    cover_elements = build_cover_elements(doc)
    insert_elements_at_start(doc, cover_elements)

    replacements = {
        "一、项目背景与方案论证": "一、绪论",
        "四、系统总体架构与类结构设计": "四、系统设计",
        "六、功能模块说明": "六、系统实现",
        "九、系统功能测试": "九、系统测试",
        "十二、总结与小组体会": "十二、总结与展望",
    }
    for old, new in replacements.items():
        replace_paragraph_text(doc, old, new)
    replace_inline_text(doc, "【待补充】", "曹丹、夏娟")
    replace_inline_text(doc, "两名同学共同完成", "两名成员共同完成")

    insert_layered_code_explanation(doc)
    insert_reference_section(doc)
    update_contents(doc)
    normalize_styles(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
