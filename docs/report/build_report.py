from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent
OUT = BASE / "高校成绩管理系统课程设计报告.docx"
UI_IMAGE = BASE / "ui-reference.png"
LOGIN_IMAGE = BASE / "frontend-login.png"
DASHBOARD_IMAGE = BASE / "frontend-dashboard.png"


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, size=11, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size, bold, color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13, True, "2E74B5" if level <= 2 else "1F4D78")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        set_font(run, 11)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(item)
        set_font(run, 11)


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_font(run, 10.5, bold)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("高校成绩管理系统课程设计报告")
    set_font(title_run, 24, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Spring Boot + Vue 3 + MySQL 8")
    set_font(subtitle_run, 13, False, "555555")

    add_para(doc, "项目类型：Web 全栈课程设计演示系统", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "实现范围：后端、前端、数据库脚本、角色权限、课程级绩点规则、成绩统计、中文报告", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_cover(doc)

    add_heading(doc, "摘要")
    add_para(
        doc,
        "高校成绩管理系统是面向高校日常教学管理场景的 Java Web 应用。系统以课程成绩为核心对象，围绕管理员、教师、学生三类用户完成基础资料维护、课程管理、绩点规则维护、成绩录入查询和统计分析。"
        "本项目后端以 Java 语言为主要实现语言，采用面向对象思想组织实体类、控制器类、业务服务类、数据访问接口和异常处理类；同时使用 Spring Boot 构建 Web 服务，使用 Spring Security 与 JWT 实现登录认证和接口权限控制，采用原生 MyBatis Mapper 接口和 XML SQL 完成数据库访问。"
        "系统重点实现了课程级绩点规则，即不同课程可以配置不同的分数到绩点转换区间，同一总评成绩在不同课程中可以得到不同绩点，符合课程设计中对 Java 程序结构、业务逻辑封装和数据库交互的综合要求。"
    )
    add_para(doc, "关键词：Java 程序设计；面向对象；Spring Boot；MyBatis；MySQL；JWT；课程级绩点规则", bold=True)

    add_heading(doc, "一、项目背景与开发意义")
    add_para(
        doc,
        "在高校教学管理过程中，成绩数据不仅用于学生课程考核，也会影响奖学金评定、学业预警、毕业审核和教学质量分析。传统手工维护成绩容易出现数据分散、统计效率低、权限边界不清晰等问题。"
        "因此，设计一个结构清晰、权限明确、可运行演示的成绩管理系统，既能体现数据库设计、后端接口设计和前端交互设计能力，也能体现对真实教学业务流程的理解。"
    )
    add_para(
        doc,
        "本系统定位为课程设计演示版，目标不是构建复杂生产系统，而是在有限范围内完整体现 Web 全栈项目的基本组成：数据库、后端服务、前端页面、角色权限、业务规则、测试验证和设计文档。"
    )

    add_heading(doc, "二、Java 技术基础")
    add_para(doc, "本课程设计属于 Java 程序设计方向，因此系统后端重点体现 Java 语言在类设计、对象封装、接口定义、异常处理、集合处理、注解使用和单元测试方面的应用。")
    add_heading(doc, "2.1 Java 语言特性应用", level=2)
    add_table(doc, ["Java 特性", "在本系统中的应用", "说明"], [
        ["类与对象", "User、Student、Teacher、Course、Grade 等实体类", "使用类描述现实业务对象，每个对象包含对应属性"],
        ["封装", "Service 层封装成绩保存、绩点计算、权限校验", "Controller 不直接处理复杂业务，业务逻辑集中在服务类中"],
        ["接口", "UserMapper、GradeMapper 等 Mapper 接口", "通过接口定义数据访问能力，XML 中提供 SQL 映射"],
        ["异常处理", "BusinessException、GlobalExceptionHandler", "业务错误通过自定义异常抛出，统一转换为 JSON 响应"],
        ["集合框架", "List<CourseGradeRule>、List<Grade>", "用于保存查询结果、规则列表和页面表格数据"],
        ["Stream API", "规则排序、区间匹配、查找目标规则", "提高集合处理代码的可读性"],
        ["BigDecimal", "成绩、学分、绩点计算", "避免 double 在小数计算中出现精度误差"],
        ["注解", "@Service、@RestController、@RequestMapping、@PreAuthorize", "声明类职责、接口路径和权限要求"],
        ["单元测试", "JUnit 5 测试 GradeRuleService", "验证规则覆盖、区间重叠、缺口和课程级绩点计算"]
    ])
    add_heading(doc, "2.2 后端技术基础", level=2)
    add_para(doc, "在 Java 语言基础上，本系统使用 Spring Boot、Spring Security、MyBatis 等框架完成 Web 项目开发。框架的作用是减少重复基础代码，使开发重点集中在业务逻辑和系统结构上。")
    add_table(doc, ["技术", "用途", "选择原因"], [
        ["Spring Boot", "构建后端 REST API", "约定优于配置，启动简单，适合课程设计快速搭建单体应用"],
        ["Spring Security + JWT", "登录认证和接口鉴权", "前后端分离场景下使用 Token 传递登录状态，便于角色权限控制"],
        ["MyBatis XML", "数据库访问", "SQL 显式写在 XML 中，便于说明表结构、查询逻辑和课程设计要求"],
        ["MySQL 8", "关系型数据库", "适合保存用户、课程、成绩等结构化数据，支持外键和索引"],
        ["Vue 3 + Vite", "前端单页应用", "开发启动快，组件化清晰，适合后台管理界面"],
        ["Element Plus", "UI 组件库", "表格、表单、弹窗、菜单组件成熟，符合后台管理系统风格"],
        ["Pinia + Axios", "状态管理和接口请求", "Pinia 保存登录状态，Axios 统一附加 JWT Token"]
    ])

    add_heading(doc, "2.3 Java 项目包结构", level=2)
    add_para(doc, "后端代码按照职责划分为多个 Java 包，既体现面向对象分层思想，也方便课程设计答辩时按模块讲解。")
    add_table(doc, ["包名", "代表类", "职责"], [
        ["controller", "AuthController、GradeController", "接收 HTTP 请求并返回统一响应"],
        ["service", "GradeService、GradeRuleService", "封装业务规则、权限校验和成绩处理逻辑"],
        ["mapper", "UserMapper、GradeMapper", "定义数据库访问接口，由 MyBatis XML 绑定 SQL"],
        ["entity", "User、Course、Grade", "描述数据库表对应的 Java 对象"],
        ["dto", "LoginRequest、GradeSaveRequest", "封装前后端传输数据"],
        ["security", "JwtAuthenticationFilter、SecurityUtils", "处理 JWT 解析、当前用户获取和权限辅助"],
        ["common", "ApiResponse、BusinessException", "统一响应结构和业务异常"]
    ])

    add_heading(doc, "三、可行性分析")
    add_heading(doc, "3.1 技术可行性", level=2)
    add_para(
        doc,
        "系统采用的 Spring Boot、Vue 3、MyBatis 和 MySQL 都是成熟且资料丰富的主流技术。项目为单体应用，不涉及分布式部署和微服务治理，因此技术复杂度适中，适合课程设计周期内完成。"
    )
    add_heading(doc, "3.2 经济可行性", level=2)
    add_para(
        doc,
        "系统使用开源框架和本地 MySQL 数据库，不需要额外购买商业软件或云服务。开发和演示均可在普通个人电脑上完成，成本较低。"
    )
    add_heading(doc, "3.3 操作可行性", level=2)
    add_para(
        doc,
        "前端采用常见后台管理系统布局，左侧为菜单，右侧为内容区。管理员、教师、学生登录后看到不同菜单，操作流程直观，便于答辩演示和教师验收。"
    )

    add_heading(doc, "四、需求分析")
    add_heading(doc, "4.1 用户角色需求", level=2)
    add_table(doc, ["角色", "核心需求", "权限边界"], [
        ["管理员", "维护用户、学生、教师、班级、课程、绩点规则和成绩数据", "可以管理所有数据"],
        ["教师", "查看本人课程，录入和修改本人课程成绩，查看课程统计", "不能编辑其他教师课程成绩"],
        ["学生", "查看个人信息、本人课程成绩和成绩汇总", "不能查看或修改其他学生数据"]
    ])
    add_heading(doc, "4.2 功能需求", level=2)
    add_bullets(doc, [
        "登录认证：用户输入用户名和密码，后端校验成功后返回 JWT Token。",
        "角色菜单：前端根据角色显示不同菜单，减少误操作。",
        "基础资料管理：管理员维护用户、学生、教师、班级和课程数据。",
        "课程绩点规则管理：管理员按课程维护分数区间、绩点值和等级标签。",
        "成绩录入与查询：教师或管理员录入成绩，系统自动计算绩点；学生只能查询本人记录。",
        "统计分析：提供系统概览、课程统计和学生个人成绩汇总。"
    ])
    add_heading(doc, "4.3 非功能需求", level=2)
    add_bullets(doc, [
        "可运行性：后端、前端和数据库脚本可以独立启动和演示。",
        "安全性：关键接口必须在后端进行角色和数据归属校验。",
        "可维护性：后端分层清晰，SQL 位于 XML 文件中，便于后期修改。",
        "易用性：界面采用后台管理系统风格，表格和表单布局清晰。"
    ])

    add_heading(doc, "五、系统总体设计")
    add_para(
        doc,
        "系统采用前后端分离结构。浏览器端通过 Vue 3 应用展示页面，Axios 将请求发送到 Spring Boot 后端。后端控制器接收请求后调用服务层，服务层完成权限校验和业务处理，再通过 MyBatis Mapper 与 MySQL 数据库交互。"
    )
    add_table(doc, ["层次", "主要组件", "职责"], [
        ["表现层", "Vue 页面、Element Plus 组件", "展示菜单、表格、表单、统计卡片和操作反馈"],
        ["接口层", "Spring Boot Controller", "接收 HTTP 请求，返回统一 JSON 结构"],
        ["业务层", "Service", "处理权限校验、成绩保存、绩点计算、统计汇总"],
        ["持久层", "Mapper 接口 + XML SQL", "执行增删改查和统计查询"],
        ["数据层", "MySQL 8", "保存系统业务数据"]
    ])
    add_para(
        doc,
        "后端统一返回结构包含 code、message、data 三个字段。前端 Axios 响应拦截器统一解析该结构，如果接口返回失败信息则显示提示，成功时直接返回 data，减少页面代码重复。"
    )

    add_heading(doc, "六、Java 后端程序设计")
    add_heading(doc, "6.1 面向对象类设计", level=2)
    add_para(
        doc,
        "系统将业务对象抽象为 Java 类。例如 User 表示登录用户，Student 表示学生档案，Teacher 表示教师档案，Course 表示课程，CourseGradeRule 表示某门课程的绩点规则，Grade 表示成绩记录。"
        "这些类与数据库表字段基本对应，便于 MyBatis 将查询结果映射为 Java 对象。"
    )
    add_table(doc, ["类名", "主要属性", "作用"], [
        ["User", "id、username、password、realName、role、status", "保存登录账号、真实姓名、角色和状态"],
        ["Student", "id、userId、classId、studentNo、name", "保存学生学号、姓名、班级和账号关联"],
        ["Teacher", "id、userId、teacherNo、name、title", "保存教师工号、姓名、职称和账号关联"],
        ["Course", "id、courseCode、courseName、credit、teacherId", "保存课程编号、名称、学分和授课教师"],
        ["CourseGradeRule", "courseId、minScore、maxScore、gradePoint、label", "保存某门课程的成绩区间和绩点映射"],
        ["Grade", "studentId、courseId、totalScore、gradePoint", "保存学生课程成绩和自动计算出的绩点"]
    ])

    add_heading(doc, "6.2 Java 分层调用流程", level=2)
    add_para(
        doc,
        "以保存成绩为例，前端请求首先进入 GradeController，Controller 负责接收 GradeSaveRequest；随后调用 GradeService，Service 层校验当前用户角色、教师课程归属和绩点规则；最后通过 GradeMapper 写入数据库。"
        "这种分层方式避免把 SQL、权限和业务计算全部写在一个类中，使 Java 程序结构更清晰。"
    )
    add_numbers(doc, [
        "Controller 层负责参数接收和接口路径定义。",
        "Service 层负责业务规则，例如教师只能修改本人课程成绩。",
        "Mapper 接口负责定义数据库操作方法。",
        "Mapper XML 负责写具体 SQL。",
        "Entity 和 DTO 负责承载数据。"
    ])

    add_heading(doc, "6.3 Java 异常与统一响应", level=2)
    add_para(
        doc,
        "系统定义了 BusinessException 表示业务异常。例如绩点规则没有覆盖 0-100、教师修改非本人课程成绩、学生尝试录入成绩等情况，都会抛出该异常。"
        "GlobalExceptionHandler 使用 @RestControllerAdvice 统一捕获异常，并返回 ApiResponse 结构，使前端能够用统一方式展示错误信息。"
    )

    add_heading(doc, "6.4 Java 集合与精确计算", level=2)
    add_para(
        doc,
        "绩点规则校验使用 List<CourseGradeRule> 保存某门课程的所有规则，并通过 Stream 排序和查找目标区间。成绩和绩点使用 BigDecimal 类型保存，避免浮点数计算导致的小数精度问题。"
        "例如 59.99、89.99、3.7 等值都适合使用 BigDecimal 表示。"
    )

    add_heading(doc, "七、功能模块详细设计")
    add_heading(doc, "7.1 登录与权限模块", level=2)
    add_para(
        doc,
        "用户登录时，前端提交用户名和密码到 /api/auth/login。后端根据 username 查询 users 表，校验密码和账号状态。校验成功后生成 JWT，并返回用户角色。前端将 Token 保存到本地，在后续请求中通过 Authorization 请求头传递。"
    )
    add_heading(doc, "7.2 基础信息管理模块", level=2)
    add_para(
        doc,
        "基础信息管理包括用户、学生、教师、班级和课程。管理员通过表格查看数据，通过弹窗表单新增或编辑记录。学生表与用户表、班级表关联，教师表与用户表关联，课程表与教师表关联。"
    )
    add_heading(doc, "7.3 绩点规则管理模块", level=2)
    add_para(
        doc,
        "绩点规则是本系统的重点模块。管理员先选择课程，再维护该课程下的多个分数区间。每条规则包含最低分、最高分、绩点和等级标签。保存规则时，后端会检查区间是否覆盖 0-100，是否存在重叠或缺口。"
    )
    add_heading(doc, "7.4 成绩管理模块", level=2)
    add_para(
        doc,
        "成绩表保存学生、课程、平时成绩、期末成绩、总评成绩、绩点和备注。管理员可以管理全部成绩，教师只能管理本人课程的成绩，学生只能查看本人课程成绩。保存成绩时，后端根据课程 ID 查询对应绩点规则，并按总评成绩计算 grade_point。"
    )
    add_heading(doc, "7.5 统计分析模块", level=2)
    add_para(
        doc,
        "统计分析包括管理员系统概览、课程统计和学生个人汇总。课程统计包含平均分、最高分、最低分、及格率和分数段分布；学生汇总包含课程数、通过课程数、平均分和平均绩点。"
    )

    add_heading(doc, "八、数据库设计")
    add_heading(doc, "8.1 数据表概览", level=2)
    add_table(doc, ["表名", "用途", "关键字段"], [
        ["users", "登录账号和角色", "id、username、password、real_name、role、status"],
        ["classes", "班级基础信息", "id、class_name、grade_year、major"],
        ["students", "学生档案", "id、user_id、class_id、student_no、name、phone"],
        ["teachers", "教师档案", "id、user_id、teacher_no、name、title"],
        ["courses", "课程信息", "id、course_code、course_name、credit、teacher_id、semester"],
        ["course_grade_rules", "课程绩点规则", "id、course_id、min_score、max_score、grade_point、label"],
        ["grades", "成绩记录", "id、student_id、course_id、total_score、grade_point"]
    ])
    add_heading(doc, "8.2 关键表字段说明", level=2)
    add_table(doc, ["字段", "所在表", "说明"], [
        ["role", "users", "取值为 ADMIN、TEACHER、STUDENT，用于角色权限控制"],
        ["teacher_id", "courses", "课程所属教师，教师权限校验依赖该字段"],
        ["course_id", "course_grade_rules", "表示规则属于哪一门课程，实现课程级绩点规则"],
        ["min_score / max_score", "course_grade_rules", "规则分数区间，要求覆盖 0-100 且不能重叠"],
        ["grade_point", "grades", "成绩保存时按课程规则计算得到的绩点"],
        ["student_id + course_id", "grades", "唯一约束，避免同一学生同一课程重复录入"]
    ])
    add_heading(doc, "8.3 示例数据设计", level=2)
    add_para(
        doc,
        "种子数据提供管理员、教师和学生账号，并内置高等数学、Java 程序设计、数据库原理等课程。高等数学和 Java 程序设计配置了不同绩点规则，便于演示同一分数在不同课程下产生不同绩点。"
    )

    add_heading(doc, "九、接口设计")
    add_table(doc, ["接口", "方法", "角色", "功能说明"], [
        ["/api/auth/login", "POST", "所有用户", "登录并获取 JWT"],
        ["/api/auth/me", "GET", "已登录用户", "获取当前登录用户信息"],
        ["/api/users", "GET/POST/PUT/DELETE", "管理员", "用户管理"],
        ["/api/classes", "GET/POST/PUT/DELETE", "管理员", "班级管理"],
        ["/api/students", "GET/POST/PUT/DELETE", "管理员/学生", "管理员管理学生，学生仅查看本人"],
        ["/api/teachers", "GET/POST/PUT/DELETE", "管理员/教师", "管理员管理教师，教师仅查看本人"],
        ["/api/courses", "GET/POST/PUT/DELETE", "管理员/教师", "管理员管理课程，教师仅查看本人课程"],
        ["/api/courses/{courseId}/grade-rules", "GET/POST", "管理员", "查询和保存课程绩点规则"],
        ["/api/grades", "GET/POST/PUT/DELETE", "管理员/教师/学生", "按角色过滤成绩数据"],
        ["/api/statistics/overview", "GET", "管理员", "系统概览统计"],
        ["/api/statistics/courses/{courseId}", "GET", "管理员/教师", "课程成绩统计"],
        ["/api/statistics/student-summary", "GET", "学生", "学生个人成绩汇总"]
    ])

    add_heading(doc, "十、关键业务实现")
    add_heading(doc, "10.1 课程级绩点计算", level=2)
    add_para(
        doc,
        "系统没有使用一张全局绩点表，而是在 course_grade_rules 表中按 course_id 保存规则。保存成绩时，GradeService 会根据 request.courseId 查询该课程规则，再调用 GradeRuleService.calculateGradePoint 计算绩点。"
    )
    add_numbers(doc, [
        "读取课程对应的所有绩点规则，并按 min_score 升序排列。",
        "检查第一条规则是否从 0 分开始，最后一条规则是否到 100 分结束。",
        "逐条检查相邻规则是否重叠，是否存在超过 0.01 分的缺口。",
        "根据 total_score 找到 min_score <= total_score <= max_score 的规则。",
        "将匹配规则的 grade_point 写入 grades.grade_point。"
    ])
    add_para(
        doc,
        "例如，种子数据中高等数学 85 分匹配 80-89.99 区间，绩点为 3.7；Java 程序设计 85 分匹配 85-94.99 区间，绩点为 3.3。这个示例可以直接用于答辩说明系统满足“课程级绩点规则”的业务要求。"
    )
    add_heading(doc, "10.2 教师课程权限控制", level=2)
    add_para(
        doc,
        "教师保存或删除成绩时，后端会根据当前登录用户找到对应 teacher_id，再检查成绩对应课程的 teacher_id 是否一致。如果课程不属于当前教师，则抛出业务异常，避免教师操作其他课程成绩。"
    )
    add_heading(doc, "10.3 学生数据隔离", level=2)
    add_para(
        doc,
        "学生查询成绩时，后端不会接收前端传入的 student_id 作为可信条件，而是根据 JWT 中的当前用户 ID 查询 students 表得到真实 student_id，再用该 ID 查询成绩，从而保证学生只能看到本人数据。"
    )

    add_heading(doc, "十一、前端界面设计")
    add_para(
        doc,
        "前端采用后台管理系统常见结构：登录页、主布局、左侧深色菜单、顶部用户栏、右侧内容区域。管理员、教师、学生登录后根据角色生成不同菜单。页面主要使用 Element Plus 的 Table、Form、Dialog、Descriptions、Alert 等组件。"
    )
    add_table(doc, ["页面", "主要组件", "说明"], [
        ["登录页", "输入框、密码框、登录按钮", "提供演示账号提示"],
        ["主布局", "侧边栏、顶部栏、内容区", "根据角色显示菜单"],
        ["管理页面", "表格、弹窗表单", "支持新增、编辑、删除"],
        ["绩点规则页", "课程选择、可编辑规则表格", "维护课程分数区间和绩点"],
        ["统计页", "统计卡片、描述列表、分布表格", "展示系统和课程统计"]
    ])

    add_heading(doc, "十二、Java 程序测试")
    add_heading(doc, "12.1 测试环境", level=2)
    add_table(doc, ["项目", "内容"], [
        ["后端运行环境", "Java 17+、Maven、Spring Boot"],
        ["前端运行环境", "Node.js、Vite"],
        ["数据库", "MySQL 8"],
        ["浏览器", "Chrome 或 Edge"]
    ])
    add_heading(doc, "12.2 测试用例", level=2)
    add_table(doc, ["编号", "测试内容", "预期结果", "结果"], [
        ["T01", "管理员账号登录", "登录成功，显示管理员菜单", "通过"],
        ["T02", "教师账号登录", "只显示我的课程、成绩录入、课程统计", "通过"],
        ["T03", "学生账号登录", "只显示个人信息、我的成绩、成绩汇总", "通过"],
        ["T04", "保存覆盖 0-100 的绩点规则", "保存成功", "通过"],
        ["T05", "保存重叠区间规则", "后端拒绝并提示区间不能重叠", "通过"],
        ["T06", "保存存在缺口的规则", "后端拒绝并提示必须覆盖 0-100", "通过"],
        ["T07", "录入高等数学 85 分", "绩点计算为 3.7", "通过"],
        ["T08", "录入 Java 程序设计 85 分", "绩点计算为 3.3", "通过"],
        ["T09", "学生查看其他学生成绩", "后端按当前学生过滤，不能越权", "通过"],
        ["T10", "教师修改非本人课程成绩", "后端拒绝操作", "通过"]
    ])
    add_heading(doc, "12.3 JUnit 单元测试说明", level=2)
    add_para(
        doc,
        "系统针对 GradeRuleService 编写了 JUnit 5 单元测试。测试先验证有效规则可以通过，再验证重叠规则和存在缺口的规则会抛出 BusinessException，最后验证高等数学和 Java 程序设计在同样 85 分时能够得到不同绩点。"
        "这些测试体现了 Java 程序设计中“先定义预期行为，再实现业务逻辑”的思想。"
    )
    add_heading(doc, "12.4 自动化验证结果", level=2)
    add_bullets(doc, [
        "后端执行 mvn test 通过，包含绩点规则覆盖、重叠、缺口和课程级计算测试。",
        "前端执行 npm run build 通过，说明 Vue 页面、路由和组件可以正常编译。",
        "DOCX 报告由脚本生成，内容与当前项目结构、数据库表和接口保持一致。"
    ])

    add_heading(doc, "十三、系统运行与部署说明")
    add_numbers(doc, [
        "安装 MySQL 8，使用 database/schema.sql 创建数据库和表结构。",
        "使用 database/seed.sql 导入管理员、教师、学生、课程、绩点规则和成绩示例数据。",
        "根据本机 MySQL 密码设置 DB_USERNAME 和 DB_PASSWORD，或修改 application.yml。",
        "进入 backend 目录执行 mvn spring-boot:run 启动后端。",
        "进入 frontend 目录执行 npm install 和 npm run dev -- --host 127.0.0.1 启动前端。",
        "浏览器访问 http://127.0.0.1:5173，使用演示账号登录。"
    ])

    add_heading(doc, "十四、系统截图")
    add_para(doc, "下图为系统后台管理界面参考图，体现左侧菜单、顶部用户区、统计卡片、成绩表格和课程绩点规则对比区域。")
    if UI_IMAGE.exists():
        doc.add_picture(str(UI_IMAGE), width=Inches(6.2))
    add_para(doc, "登录页面截图：")
    if LOGIN_IMAGE.exists():
        doc.add_picture(str(LOGIN_IMAGE), width=Inches(6.2))
    add_para(doc, "管理员控制台截图：")
    if DASHBOARD_IMAGE.exists():
        doc.add_picture(str(DASHBOARD_IMAGE), width=Inches(6.2))

    add_heading(doc, "十五、总结与心得")
    add_para(
        doc,
        "本项目完成了高校成绩管理系统课程设计要求的核心内容，包含可运行 Java 后端、可运行前端、MySQL 建表与种子数据、角色登录权限、课程级绩点规则、成绩录入查询、基础统计和中文 Word 报告。"
        "通过本项目，可以完整展示 Java 程序设计中的类与对象、接口、集合、异常处理、注解、单元测试，以及 Web 全栈开发中的数据库建模、接口设计、权限控制和前端页面组织。"
    )
    add_para(
        doc,
        "项目中最关键的收获是业务规则不能只停留在界面层。教师课程归属、学生数据隔离、绩点规则校验和绩点计算都必须在后端服务层完成，前端菜单隐藏只能作为用户体验优化，不能作为真正的安全边界。"
    )
    add_para(
        doc,
        "后续如果继续扩展，可以增加 Excel 导入导出、成绩审批、操作日志、更多图表统计等功能。但在本课程设计中，系统保持了标准演示版范围，避免了与核心成绩管理无关的大型功能。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
