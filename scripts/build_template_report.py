from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path("D:/15Pro/Desktop")
SOURCE_REPORT = DESKTOP / "高校成绩管理系统课程设计报告.docx"
OUTPUT = DESKTOP / "高校成绩管理系统课程设计报告_按模板重排.docx"
MEDIA_DIR = DESKTOP / "campus_report_media"


TOC_ENTRIES = [
    ("1 绪论", "1"),
    ("1.1 课题研究背景", "1"),
    ("1.2 课题目的及意义", "1"),
    ("1.3 开发环境与技术路线", "2"),
    ("2 需求分析", "2"),
    ("2.1 用户角色需求", "2"),
    ("2.2 功能需求", "2"),
    ("2.3 非功能需求", "3"),
    ("3 系统设计", "3"),
    ("3.1 总体架构设计", "3"),
    ("3.2 功能模块设计", "3"),
    ("3.3 数据库设计", "3"),
    ("3.4 核心业务规则设计", "4"),
    ("4 系统实现", "4"),
    ("4.1 后端关键实现", "4"),
    ("4.2 前端关键实现", "5"),
    ("4.3 系统运行界面", "5"),
    ("5 系统测试", "6"),
    ("5.1 测试环境", "6"),
    ("5.2 测试用例及结果", "7"),
    ("6 总结与展望", "7"),
    ("参考文献", "8"),
    ("附录A 主要接口清单", "8"),
]


def ensure_media() -> list[Path]:
    MEDIA_DIR.mkdir(exist_ok=True)
    if SOURCE_REPORT.exists():
        with zipfile.ZipFile(SOURCE_REPORT) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/"):
                    target = MEDIA_DIR / Path(name).name
                    if not target.exists():
                        target.write_bytes(zf.read(name))
    return sorted(MEDIA_DIR.glob("image*.png"))


def set_run_font(run, name="宋体", size=12, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman" if name == "宋体" else name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman" if name == "宋体" else name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, *, align=None, first_line=True, before=0, after=0, line=18):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(line)
    if first_line:
        fmt.first_line_indent = Pt(24)
    else:
        fmt.first_line_indent = Pt(0)
    if align is not None:
        paragraph.alignment = align


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(str(text))
    set_run_font(run, "宋体", size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="666666", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_paragraph(doc, text="", *, style=None, bold=False, size=12, align=None, first_line=True, before=0, after=0):
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, align=align, first_line=first_line, before=before, after=after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph(style="Heading 1")
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=12, after=6, line=20)
        run = p.add_run(text)
        set_run_font(run, "黑体", 16, True)
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=8, after=4, line=18)
        run = p.add_run(text)
        set_run_font(run, "黑体", 14, True)
    else:
        p = doc.add_paragraph(style="Heading 3")
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=6, after=3, line=18)
        run = p.add_run(text)
        set_run_font(run, "黑体", 12, True)
    return p


def add_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=10.5)
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], text, size=10.5, align=align)
    if widths_cm:
        set_table_width(table, widths_cm)
    add_paragraph(doc, "", first_line=False, after=3)
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="BFBFBF", size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, "F5F5F5")
    cell.text = ""
    for line in code.strip().splitlines():
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        r = p.add_run(line.rstrip())
        set_run_font(r, "Courier New", 8.5)
    if cell.paragraphs and not cell.paragraphs[0].text:
        cell._tc.remove(cell.paragraphs[0]._p)
    add_paragraph(doc, "", first_line=False, after=2)


def add_caption(doc, text):
    p = add_paragraph(doc, text, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=3, after=6)
    for run in p.runs:
        set_run_font(run, "宋体", 10.5)


def add_picture_with_caption(doc, img_path, caption):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.line_spacing = 1
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.7))
        add_caption(doc, caption)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("高校成绩管理系统课程设计报告")
    set_run_font(run, "宋体", 9)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = Pt(18)
    normal.paragraph_format.space_after = Pt(0)

    for name, font, size, bold in [
        ("Heading 1", "黑体", 16, True),
        ("Heading 2", "黑体", 14, True),
        ("Heading 3", "黑体", 12, True),
    ]:
        style = styles[name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_cover(doc):
    for _ in range(3):
        add_paragraph(doc, "", first_line=False)
    p = add_paragraph(doc, "课  程  设  计", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=0, after=28)
    set_run_font(p.runs[0], "黑体", 28, True)
    p = add_paragraph(doc, "高校成绩管理系统", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=0, after=60)
    set_run_font(p.runs[0], "黑体", 18, True)

    rows = [
        ("课程设计名称：", "Java课程设计"),
        ("专 业 班 级 ：", "物联网2403"),
        ("学 生 姓 名 ：", "周成林    逯雨喆"),
        ("学       号 ：", "241040700321    241040700308"),
        ("指 导 教 师 ：", "【待补充】"),
        ("课程设计时间 ：", "2026.6.8—2026.6.20"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="000000", size="8")
    set_table_width(table, [4.8, 9.4])
    for r, (k, v) in enumerate(rows):
        set_cell_text(table.cell(r, 0), k, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(r, 1), v, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
        table.rows[r].height = Cm(1.0)
    doc.add_page_break()


def add_toc(doc):
    p = add_paragraph(doc, "目   次", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=0, after=18)
    set_run_font(p.runs[0], "黑体", 15, True)
    for title, page in TOC_ENTRIES:
        indent = Pt(24) if "." in title[:5] else Pt(0)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        dots = "…" * max(4, 26 - len(title))
        r = p.add_run(f"{title}{dots}{page}")
        set_run_font(r, "宋体", 12)
    doc.add_page_break()


def add_task_info(doc):
    add_heading(doc, "课程设计任务书信息", 1)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ("课程名称", "Java程序设计课程设计"),
            ("设计题目", "高校成绩管理系统"),
            ("项目类型", "Web 全栈管理系统演示版"),
            ("成员姓名", "周成林、逯雨喆"),
            ("班级", "物联网2403"),
            ("成员学号", "周成林：241040700321；逯雨喆：241040700308"),
            ("主要目的", "巩固 Java 面向对象编程、数据库访问、Web 后端接口设计和前后端联调能力，完成可运行的信息管理系统。"),
            ("功能要求", "登录权限、用户管理、学生教师班级课程管理、课程成绩规则管理、成绩录入查询、统计分析。"),
        ],
        [4.0, 12.2],
    )
    add_paragraph(
        doc,
        "本报告依据《Java课程设计任务书》和课程设计报告模板重新整理，重点说明高校成绩管理系统的需求分析、系统设计、数据库设计、核心实现、测试结果和总结。",
    )


def add_body(doc, images):
    add_task_info(doc)

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 课题研究背景", 2)
    add_paragraph(doc, "高校教学管理中，成绩数据直接服务于课程考核、奖学金评定、学业预警和教学质量分析。传统纸质记录或电子表格方式存在数据分散、重复维护、统计效率低和权限边界不清晰等问题，难以满足多角色协同管理需要。")
    add_paragraph(doc, "本课题设计并实现“高校成绩管理系统”，通过浏览器访问后台管理界面，以 Spring Boot 后端、MySQL 数据库和 Vue 3 前端组成完整的 Web 全栈系统。系统面向管理员、教师和学生三类用户，完成成绩数据从录入、计算、查询到统计的闭环管理。")
    add_heading(doc, "1.2 课题目的及意义", 2)
    add_paragraph(doc, "本课程设计的目标不是单纯展示界面，而是完成一个可运行、可演示、可答辩的 Java Web 项目。通过项目开发，可以综合训练 Java 面向对象设计、分层架构、接口封装、异常处理、MyBatis XML SQL、JWT 权限控制、前后端接口联调和系统测试等能力。")
    add_paragraph(doc, "系统的业务重点是课程级成绩计算规则。不同课程可以单独设置平时成绩权重、期末成绩权重和绩点换算区间，录入成绩时只输入平时成绩和期末成绩，由后端计算总评成绩和绩点，避免把所有课程套用同一套全局规则。")
    doc.add_page_break()
    add_heading(doc, "1.3 开发环境与技术路线", 2)
    add_table(
        doc,
        ["层次", "技术", "作用"],
        [
            ("前端", "Vue 3、Vite、Element Plus", "构建后台管理系统风格页面，提供表格、表单、弹窗和菜单导航。"),
            ("状态与接口", "Pinia、Axios", "保存登录用户状态，集中处理接口请求与 JWT Token。"),
            ("后端", "Spring Boot、Spring Security、JWT", "提供 REST API、登录认证、接口鉴权和统一异常处理。"),
            ("持久层", "原生 MyBatis Mapper 接口 + XML SQL", "显式编写 SQL，便于说明数据库访问过程。"),
            ("数据库", "MySQL 8", "保存用户、班级、学生、教师、课程、成绩规则和成绩数据。"),
        ],
        [3.0, 5.0, 8.2],
    )
    add_paragraph(doc, "系统采用前后端分离模式：前端通过 Axios 请求 /api 接口，后端 Controller 接收请求后调用 Service 完成业务处理，再通过 Mapper 和 XML SQL 访问 MySQL。")

    add_heading(doc, "2 需求分析", 1)
    add_heading(doc, "2.1 用户角色需求", 2)
    add_table(
        doc,
        ["角色", "核心需求", "权限边界"],
        [
            ("管理员", "维护用户、学生、教师、班级、课程、成绩规则、成绩数据，并查看统计信息。", "可以管理系统内全部业务数据。"),
            ("教师", "查看本人课程，配置本人课程成绩规则，录入和修改本人课程成绩，查看课程统计。", "不能管理其他教师课程、成绩和规则。"),
            ("学生", "查看个人信息、本人课程成绩和成绩汇总。", "不能录入成绩，不能查看其他学生成绩。"),
        ],
        [2.4, 7.1, 6.7],
    )
    add_heading(doc, "2.2 功能需求", 2)
    add_table(
        doc,
        ["模块", "功能说明"],
        [
            ("登录与权限", "用户输入用户名和密码登录，后端验证成功后返回 JWT，前端按角色显示菜单。"),
            ("用户管理", "管理员维护系统账号、姓名、角色和状态。"),
            ("学生管理", "管理员维护学生学号、姓名、班级、联系方式等档案。"),
            ("教师管理", "管理员维护教师工号、姓名、职称和联系方式。"),
            ("班级管理", "管理员维护班级名称、年级和专业信息。"),
            ("课程管理", "管理员维护课程编码、课程名称、学分、任课教师、学期和成绩权重。"),
            ("成绩规则管理", "管理员管理所有课程规则，教师只能管理本人课程规则。"),
            ("成绩录入查询", "管理员或教师录入平时成绩和期末成绩，后端自动计算总评成绩和绩点。"),
            ("统计分析", "提供系统概览、课程统计和学生个人成绩汇总。"),
        ],
        [3.2, 13.0],
    )
    add_heading(doc, "2.3 非功能需求", 2)
    add_paragraph(doc, "安全性方面，系统需要基于 JWT 识别当前登录用户，并在后端服务层执行角色权限判断，保证学生只能查看本人数据、教师只能操作本人课程。可靠性方面，课程成绩规则必须经过后端校验后才允许成绩录入，避免错误数据进入成绩表。可维护性方面，后端按 controller、service、mapper、entity、dto、security 等包组织代码，前端按 layout、router、stores、views、utils 等目录组织页面。")

    add_heading(doc, "3 系统设计", 1)
    add_heading(doc, "3.1 总体架构设计", 2)
    add_table(
        doc,
        ["层次", "主要组件", "职责"],
        [
            ("表现层", "Vue 页面、Element Plus 组件", "展示登录页、侧边菜单、表格、表单、弹窗和统计结果。"),
            ("接口层", "Spring Boot Controller", "接收 HTTP 请求，返回统一 ApiResponse。"),
            ("业务层", "Service", "处理成绩计算、规则校验、角色权限和统计逻辑。"),
            ("持久层", "Mapper 接口 + XML SQL", "执行增删改查和统计查询。"),
            ("数据层", "MySQL 8", "保存系统业务数据。"),
        ],
        [2.6, 5.7, 7.9],
    )
    add_heading(doc, "3.2 功能模块设计", 2)
    add_table(
        doc,
        ["包名", "代表类", "职责"],
        [
            ("controller", "AuthController、GradeController、GradeRuleController", "定义 REST API 入口。"),
            ("service", "GradeService、GradeRuleService、CourseService", "封装核心业务逻辑和权限判断。"),
            ("mapper", "UserMapper、GradeMapper、CourseGradeRuleMapper", "定义数据库访问接口。"),
            ("entity", "User、Student、Teacher、Course、Grade", "承载数据库表对应的数据对象。"),
            ("dto", "LoginRequest、GradeSaveRequest、CourseGradeRuleConfig", "封装前后端传输数据。"),
            ("security", "JwtTokenProvider、JwtAuthenticationFilter、SecurityUtils", "处理 JWT、当前用户和角色判断。"),
            ("common", "ApiResponse、BusinessException、GlobalExceptionHandler", "统一响应结构和异常处理。"),
        ],
        [2.8, 6.7, 6.7],
    )
    add_heading(doc, "3.3 数据库设计", 2)
    add_table(
        doc,
        ["表名", "用途", "关键字段"],
        [
            ("users", "登录账号和角色", "id、username、password、real_name、role、status"),
            ("classes", "班级信息", "id、class_name、grade_year、major"),
            ("students", "学生档案", "id、user_id、class_id、student_no、name"),
            ("teachers", "教师档案", "id、user_id、teacher_no、name、title"),
            ("courses", "课程信息与权重配置", "id、course_code、course_name、teacher_id、usual_weight、final_weight"),
            ("course_grade_rules", "课程绩点换算规则", "id、course_id、min_score、max_score、grade_point、label"),
            ("grades", "成绩记录", "id、student_id、course_id、usual_score、final_score、total_score、grade_point"),
        ],
        [3.8, 4.5, 7.9],
    )
    add_paragraph(doc, "courses 表中的 usual_weight 和 final_weight 保存每门课程自己的平时成绩和期末成绩占比；course_grade_rules 表通过 course_id 与课程表关联，实现课程级绩点换算；grades 表保存用户输入分数和后端计算后的 total_score、grade_point。")
    add_heading(doc, "3.4 核心业务规则设计", 2)
    add_table(
        doc,
        ["规则", "实现要求"],
        [
            ("课程权重独立", "courses 表保存 usual_weight 和 final_weight，不使用全局权重。"),
            ("权重校验", "usual_weight + final_weight 必须等于 100，且权重范围在 0-100 之间。"),
            ("总评计算", "total_score = usual_score * usual_weight / 100 + final_score * final_weight / 100。"),
            ("绩点计算", "根据 course_grade_rules 中对应 course_id 的分数区间匹配 grade_point。"),
            ("规则完整性", "同一课程的绩点区间必须覆盖 0-100，不能重叠，不能存在缺口。"),
            ("权限控制", "教师只能管理本人课程的成绩和规则，学生只能查看本人数据。"),
        ],
        [4.0, 12.2],
    )

    add_heading(doc, "4 系统实现", 1)
    add_heading(doc, "4.1 后端关键实现", 2)
    add_paragraph(doc, "后端核心实现集中在 GradeRuleService、GradeService、CourseService 等业务类中。成绩录入时，前端请求中不传 total_score 和 grade_point，后端读取课程权重和课程绩点规则后统一计算，保证业务结果可信。")
    add_code_block(
        doc,
        """
public BigDecimal calculateTotalScore(BigDecimal usualScore, BigDecimal finalScore,
                                       BigDecimal usualWeight, BigDecimal finalWeight) {
    validateScore(usualScore, "平时成绩");
    validateScore(finalScore, "期末成绩");
    validateWeights(usualWeight, finalWeight);
    return usualScore.multiply(usualWeight)
            .add(finalScore.multiply(finalWeight))
            .divide(new BigDecimal("100"), 2, RoundingMode.HALF_UP);
}
""",
    )
    add_code_block(
        doc,
        """
public Grade save(GradeSaveRequest request) {
    if (SecurityUtils.isStudent()) {
        throw new BusinessException("学生不能录入或修改成绩");
    }
    courseService.assertTeacherOwnsCourse(request.courseId);
    Course course = courseService.findRequiredCourse(request.courseId);
    List<CourseGradeRule> rules = ruleMapper.findByCourseId(request.courseId);
    BigDecimal totalScore = gradeRuleService.calculateTotalScore(
            request.usualScore, request.finalScore,
            course.usualWeight, course.finalWeight);
    BigDecimal gradePoint = gradeRuleService.calculateGradePoint(rules, totalScore);
    // 保存时写入后端计算结果，避免前端伪造总评成绩和绩点。
}
""",
    )
    add_paragraph(doc, "课程规则管理保存时会先校验权重是否相加为 100，再检查绩点区间是否从 0 开始、到 100 结束、相邻区间无重叠且无缺口。教师用户还需要通过 CourseService 判断课程归属。")
    add_heading(doc, "4.2 前端关键实现", 2)
    add_paragraph(doc, "前端使用 Vue Router 维护角色菜单，使用 Pinia 保存登录信息，使用 Axios 请求拦截器自动携带 JWT。管理员、教师、学生登录后看到不同页面，未登录用户会被路由守卫重定向到登录页。")
    add_code_block(
        doc,
        """
request.interceptors.request.use((config) => {
  const token = localStorage.getItem('token')
  if (token) config.headers.Authorization = `Bearer ${token}`
  return config
})

router.beforeEach((to) => {
  const auth = useAuthStore()
  if (to.path !== '/login' && !auth.isLoggedIn) return '/login'
  const roles = to.meta?.roles
  if (roles && !roles.includes(auth.role)) return '/dashboard'
  return true
})
""",
    )
    add_heading(doc, "4.3 系统运行界面", 2)
    captions = ["图4-1 登录页面", "图4-2 管理员控制台", "图4-3 成绩与规则管理界面"]
    for img, caption in zip(images, captions):
        add_picture_with_caption(doc, img, caption)

    add_heading(doc, "5 系统测试", 1)
    add_heading(doc, "5.1 测试环境", 2)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ("后端环境", "Java 17、Maven、Spring Boot"),
            ("前端环境", "Node.js、Vite、Vue 3"),
            ("数据库", "MySQL 8，数据库名 campus_grade_management"),
            ("浏览器", "Edge 或 Chrome"),
            ("测试账号", "admin/admin123，teacher01/123456，student01/123456"),
        ],
        [4.0, 12.2],
    )
    add_heading(doc, "5.2 测试用例及结果", 2)
    add_table(
        doc,
        ["编号", "测试内容", "预期结果", "结果"],
        [
            ("T01", "管理员账号登录", "登录成功，显示用户、学生、教师、班级、课程、规则、成绩、统计菜单。", "通过"),
            ("T02", "教师账号登录", "只显示我的课程、成绩规则、成绩录入、课程统计等教师菜单。", "通过"),
            ("T03", "学生账号登录", "只显示个人信息、我的成绩、成绩汇总。", "通过"),
            ("T04", "保存权重 30/70 的课程规则", "权重合法，保存成功。", "通过"),
            ("T05", "保存权重 40/50 的课程规则", "后端拒绝，提示权重之和必须为 100。", "通过"),
            ("T06", "保存重叠绩点区间", "后端拒绝，提示规则区间不能重叠。", "通过"),
            ("T07", "保存存在缺口的绩点区间", "后端拒绝，提示必须覆盖 0-100。", "通过"),
            ("T08", "录入成绩只填写平时和期末", "后端自动计算 total_score 和 grade_point。", "通过"),
            ("T09", "教师编辑非本人课程成绩", "后端拒绝操作。", "通过"),
            ("T10", "学生查看其他学生成绩", "接口不返回其他学生成绩。", "通过"),
            ("T11", "前端项目构建", "npm run build 编译通过。", "通过"),
            ("T12", "后端规则单元测试", "GradeRuleServiceTest 覆盖规则校验和计算场景。", "通过"),
        ],
        [1.4, 4.2, 8.8, 1.8],
    )
    add_heading(doc, "5.3 运行与部署说明", 2)
    for idx, item in enumerate(
        [
            "安装 MySQL 8，执行 database/schema.sql 创建数据库和表结构。",
            "执行 database/seed.sql 导入演示账号、基础资料、课程规则和成绩示例数据。",
            "根据本机数据库账号修改 backend/src/main/resources/application.yml，或设置 DB_USERNAME、DB_PASSWORD 环境变量。",
            "进入 backend 目录执行 mvn spring-boot:run 启动后端，默认地址为 http://localhost:8080。",
            "进入 frontend 目录执行 npm install，再执行 npm run dev -- --host 127.0.0.1 启动前端。",
            "浏览器访问 http://127.0.0.1:5173，使用演示账号登录并按角色进行功能验证。",
        ],
        1,
    ):
        add_paragraph(doc, f"{idx}. {item}", first_line=False)

    add_heading(doc, "6 总结与展望", 1)
    add_paragraph(doc, "本项目完成了高校成绩管理系统课程设计要求的核心内容，包括可运行后端、可运行前端、MySQL 建表与种子数据、JWT 登录认证、三类角色权限控制、课程级成绩计算规则、成绩录入查询、基础统计分析和中文课程设计报告。系统功能范围保持在课程设计演示版内，没有加入移动端、选课审批、复杂日志、Excel 导入导出等与核心目标无关的大型功能。")
    add_paragraph(doc, "通过本次课程设计，小组成员对 Java Web 项目的完整开发流程有了更系统的理解。尤其是在成绩规则模块中，业务规则不能只停留在前端界面提示，而必须在后端服务层进行严格校验；教师课程归属、学生数据隔离、规则覆盖校验和成绩自动计算都应由后端保证。")
    add_paragraph(doc, "后续如果继续完善系统，可以在保持当前核心架构的基础上增加成绩导出、操作日志、课程选课、成绩审批等扩展功能。但在本次课程设计范围内，系统已经能够满足演示、答辩和基础教学管理场景。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] 耿祥义. Java2实用教程（第6版）[M]. 北京：清华大学出版社, 2017.",
        "[2] Craig Walls. Spring Boot实战[M]. 北京：人民邮电出版社, 2016.",
        "[3] MyBatis 官方文档. MyBatis 3 User Guide[EB/OL].",
        "[4] Vue.js 官方文档. Vue 3 Guide[EB/OL].",
        "[5] Element Plus 官方文档. Element Plus Component Documentation[EB/OL].",
        "[6] MySQL 官方文档. MySQL 8.0 Reference Manual[EB/OL].",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_line=False)

    add_heading(doc, "附录A 主要接口清单", 1)
    add_table(
        doc,
        ["接口", "方法", "角色", "功能说明"],
        [
            ("/api/auth/login", "POST", "所有用户", "登录并获取 JWT。"),
            ("/api/auth/me", "GET", "已登录用户", "获取当前用户信息。"),
            ("/api/users", "GET/POST/PUT/DELETE", "管理员", "用户管理。"),
            ("/api/students", "GET/POST/PUT/DELETE", "管理员/学生", "管理员维护学生，学生查看本人信息。"),
            ("/api/teachers", "GET/POST/PUT/DELETE", "管理员/教师", "管理员维护教师，教师查看本人信息。"),
            ("/api/classes", "GET/POST/PUT/DELETE", "管理员", "班级管理。"),
            ("/api/courses", "GET/POST/PUT/DELETE", "管理员/教师", "管理员管理课程，教师查看本人课程。"),
            ("/api/course-grade-rules", "GET/POST", "管理员/教师", "课程成绩规则查询和保存。"),
            ("/api/grades", "GET/POST/PUT/DELETE", "管理员/教师/学生", "成绩列表、录入、修改、删除和学生查询。"),
            ("/api/statistics/overview", "GET", "管理员", "系统概览统计。"),
            ("/api/statistics/course/{id}", "GET", "管理员/教师", "课程成绩统计。"),
            ("/api/statistics/student-summary", "GET", "学生", "学生个人成绩汇总。"),
        ],
        [4.2, 3.1, 3.0, 5.9],
    )


def build():
    images = ensure_media()
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_toc(doc)
    add_body(doc, images)
    doc.core_properties.title = "高校成绩管理系统课程设计报告"
    doc.core_properties.author = "周成林、逯雨喆"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
